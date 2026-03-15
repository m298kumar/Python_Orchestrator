"""
Generate technical and business documentation Word documents
for the AI Test Case Orchestrator project.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colours ──────────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1F, 0x4E, 0x79)
BLUE       = RGBColor(0x2E, 0x75, 0xB6)
TEAL       = RGBColor(0x00, 0xB0, 0xA0)
ACCENT     = RGBColor(0x00, 0x70, 0xC0)
GREEN      = RGBColor(0x37, 0x56, 0x23)
AMBER      = RGBColor(0xC5, 0x5A, 0x11)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK       = RGBColor(0x1A, 0x1A, 0x1A)
MID        = RGBColor(0x44, 0x44, 0x44)
LIGHT      = RGBColor(0x66, 0x66, 0x66)
RED        = RGBColor(0xC0, 0x00, 0x00)

HEX_NAVY   = "1F4E79"
HEX_BLUE   = "2E75B6"
HEX_TEAL   = "00B0A0"
HEX_WHITE  = "FFFFFF"
HEX_HDR_BG = "1F4E79"
HEX_ROW_BG = "DEEAF1"
HEX_ALT    = "F2F7FB"
HEX_CODE   = "F5F5F5"
HEX_BORDER = "BFBFBF"
HEX_DARK   = "1A1A1A"
HEX_GREEN_BG = "E2EFDA"
HEX_BLUE_BG  = "DEEAF1"
HEX_LIGHT_BG = "EEF4FB"

# ── XML helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)

def set_cell_borders(cell, color_hex=HEX_BORDER, size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color_hex)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def set_table_width(table, width_twips):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblW'))
    if existing is not None:
        tblPr.remove(existing)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def set_col_width(cell, width_twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        tcPr.remove(existing)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.insert(0, tcW)

def add_row_bottom_border(paragraph, color_hex, size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_page_size(section, width_twips=12240, height_twips=15840):
    sectPr = section._sectPr
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is None:
        pgSz = OxmlElement('w:pgSz')
        sectPr.append(pgSz)
    pgSz.set(qn('w:w'), str(width_twips))
    pgSz.set(qn('w:h'), str(height_twips))

def set_page_margins(section, top=1440, right=1080, bottom=1440, left=1080):
    sectPr = section._sectPr
    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is None:
        pgMar = OxmlElement('w:pgMar')
        sectPr.append(pgMar)
    pgMar.set(qn('w:top'),    str(top))
    pgMar.set(qn('w:right'),  str(right))
    pgMar.set(qn('w:bottom'), str(bottom))
    pgMar.set(qn('w:left'),   str(left))
    pgMar.set(qn('w:header'), '720')
    pgMar.set(qn('w:footer'), '720')

# ── Document builder helpers ──────────────────────────────────────────────────
def add_heading(doc, text, level=1, color=None, space_before=240, space_after=120, bottom_border=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Twips(space_before)
    p.paragraph_format.space_after  = Twips(space_after)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = color
    if bottom_border:
        add_row_bottom_border(p, bottom_border, size=6)
    return p

def add_body(doc, text, color=DARK, bold=False, italic=False, size=11,
             space_before=60, space_after=100, font="Arial"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Twips(space_before)
    p.paragraph_format.space_after  = Twips(space_after)
    run = p.add_run(text)
    run.font.name  = font
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.bold   = bold
    run.italic = italic
    return p

def add_bullet(doc, text, level=0, color=DARK, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Twips(40)
    p.paragraph_format.space_after  = Twips(40)
    p.paragraph_format.left_indent  = Inches(0.25 + 0.25 * level)
    run = p.add_run(text)
    run.font.name  = "Arial"
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    return p

def add_numbered(doc, text, color=DARK, size=11):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Twips(60)
    p.paragraph_format.space_after  = Twips(60)
    p.paragraph_format.left_indent  = Inches(0.25)
    run = p.add_run(text)
    run.font.name  = "Arial"
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    return p

def add_code(doc, text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Twips(40)
    p.paragraph_format.space_after  = Twips(40)
    p.paragraph_format.left_indent  = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), HEX_CODE)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(size)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_gap(doc, size=100):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Twips(size)
    p.paragraph_format.space_after  = Twips(0)
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.oxml.ns.qn('w:br'))
    # Use python-docx's proper page break
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    p.runs[0]._r.append(br) if p.runs else p.add_run()._r.append(br)

def page_break(doc):
    doc.add_page_break()

# ── Table builder ─────────────────────────────────────────────────────────────
def make_table(doc, headers, rows, col_widths_twips, header_bg=HEX_HDR_BG, alt_row_bg=HEX_ALT):
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, sum(col_widths_twips))

    # Header row
    hdr_row = table.rows[0]
    for ci, (h, w) in enumerate(zip(headers, col_widths_twips)):
        cell = hdr_row.cells[ci]
        set_col_width(cell, w)
        set_cell_bg(cell, header_bg)
        set_cell_borders(cell, HEX_BLUE, 6)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.font.name  = "Arial"
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = HEX_WHITE if ri % 2 == 0 else alt_row_bg
        for ci, (text, w) in enumerate(zip(row_data, col_widths_twips)):
            cell = row.cells[ci]
            set_col_width(cell, w)
            set_cell_bg(cell, fill)
            set_cell_borders(cell, HEX_BORDER, 4)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(text)
            run.font.name  = "Arial"
            run.font.size  = Pt(10)
            run.font.bold  = (ci == 0)
            run.font.color.rgb = DARK

    doc.add_paragraph()  # spacing after table
    return table

HEX_WHITE = "FFFFFF"

# ── Info / callout box ────────────────────────────────────────────────────────
def add_callout(doc, title, lines, fill=HEX_BLUE_BG, title_color=BLUE, border_color=HEX_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, 9360)
    cell = table.rows[0].cells[0]
    set_col_width(cell, 9360)
    set_cell_bg(cell, fill)
    set_cell_borders(cell, border_color, 6)
    set_cell_margins(cell, top=120, bottom=120, left=200, right=200)

    # Title
    tp = cell.paragraphs[0]
    tp.clear()
    tr = tp.add_run(title)
    tr.font.name  = "Arial"
    tr.font.size  = Pt(11)
    tr.font.bold  = True
    tr.font.color.rgb = title_color
    tp.paragraph_format.space_after = Twips(80)

    for line in lines:
        lp = cell.add_paragraph()
        lp.paragraph_format.left_indent = Inches(0.2)
        lp.paragraph_format.space_before = Twips(30)
        lp.paragraph_format.space_after  = Twips(30)
        lr = lp.add_run(line)
        lr.font.name  = "Arial"
        lr.font.size  = Pt(10)
        lr.font.color.rgb = DARK

    doc.add_paragraph()
    return table

# ─────────────────────────────────────────────────────────────────────────────
# TECHNICAL DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────
import docx

def build_technical_doc():
    doc = Document()
    section = doc.sections[0]
    set_page_size(section)
    set_page_margins(section)

    # Default style
    style = doc.styles['Normal']
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ── Cover page ────────────────────────────────────────────────────────────
    add_gap(doc, 2000)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Twips(120)
    r = p.add_run("AI Test Case Orchestrator")
    r.font.name = "Arial"; r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Twips(80)
    r2 = p2.add_run("Technical Architecture & Developer Reference")
    r2.font.name = "Arial"; r2.font.size = Pt(16); r2.font.color.rgb = BLUE

    add_row_bottom_border(p2, HEX_TEAL, 8)

    add_gap(doc, 80)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Version 10  |  March 2026  |  For: Engineering, QA & DevOps Teams")
    r3.font.name = "Arial"; r3.font.size = Pt(10); r3.font.color.rgb = MID

    doc.add_page_break()

    # ── Section 1: Overview ───────────────────────────────────────────────────
    add_heading(doc, "1. System Overview", 1, NAVY, bottom_border=HEX_TEAL)
    add_body(doc,
        "The AI Test Case Orchestrator is a fully local, LLM-powered pipeline that converts "
        "software requirements files into structured, executable test cases. It uses a locally "
        "running Ollama LLM, ChromaDB for semantic retrieval, and an MCP (Model Context Protocol) "
        "server for integration with Claude Desktop and CI pipelines.")

    add_gap(doc, 80)
    add_heading(doc, "1.1  High-Level Pipeline", 2, BLUE)
    add_body(doc, "The five-stage pipeline runs sequentially for each invocation:")
    add_gap(doc, 60)

    make_table(doc,
        ["Stage", "What Happens"],
        [
            ["1. Parse Requirements", "Multi-format reader ingests .csv, .txt, .pdf, .docx, .xlsx, .json and normalises into Requirement objects."],
            ["2. Embed & Index",      "Requirements are embedded (Ollama / SentenceTransformer) and stored in ChromaDB. Domain vocabulary (screen names, UI elements) is extracted."],
            ["3. Contextualise",      "For each requirement the system retrieves: similar requirements, approved few-shot TC examples, and domain vocab terms."],
            ["4. Generate",          "The Ollama LLM generates test cases per acceptance criterion using type-specific prompts and 2-attempt retry logic."],
            ["5. Export",            "Output is written as CSV (standard), Zephyr Scale CSV, and a JSON metadata report."],
        ],
        [2160, 7200],
    )

    add_heading(doc, "1.2  File Structure", 2, BLUE)
    make_table(doc,
        ["File", "Responsibility"],
        [
            ["orchestrator.py",          "Main CLI entry point. Wires all components together for a full run."],
            ["config.py",                "Pydantic settings loaded from .env. Singleton config object used everywhere."],
            ["requirements_reader.py",   "Multi-format parser. Produces Requirement[] from any supported file type."],
            ["llm_client.py",            "Ollama HTTP client. Handles JSON schema enforcement, retry, and fallback."],
            ["chroma_store.py",          "ChromaDB manager. Owns three collections and the embedding function."],
            ["test_generator.py",        "AC classifier, prompt builder, sanitiser, and TestCase object factory."],
            ["mcp_server.py",            "MCP tool provider (FastMCP). Also exposes all tools as a Click CLI."],
            ["exporters/exporters.py",   "Three exporters: CSVExporter, ZephyrScaleExporter, JSONReportExporter."],
            ["environment.py",           "Behave BDD hooks (setup / teardown) for step-based test execution."],
            ["test_generation_steps.py", "Behave step definitions for BDD-style orchestration."],
        ],
        [2520, 6840],
    )

    doc.add_page_break()

    # ── Section 2: Configuration ──────────────────────────────────────────────
    add_heading(doc, "2. Configuration Reference", 1, NAVY, bottom_border=HEX_TEAL)
    add_body(doc,
        "All runtime settings are read from a .env file in the working directory. "
        "The config.py module exposes them as a typed singleton (config) using Pydantic BaseSettings.")
    add_gap(doc, 80)

    add_heading(doc, "2.1  Ollama / LLM Settings", 2, BLUE)
    make_table(doc,
        ["Variable", "Default", "Purpose"],
        [
            ["OLLAMA_BASE_URL",    "http://localhost:11434", "Ollama server endpoint. Change if running on a remote host."],
            ["OLLAMA_MODEL",       "phi4-mini:3.8b",        "Model to use. phi4-mini: 2.4 GB, ~25 tok/s on CPU."],
            ["OLLAMA_TEMPERATURE", "0.6",                   "Sampling temperature. Lower = more deterministic output."],
            ["OLLAMA_NUM_CTX",     "8192",                  "Context window in tokens. Must fit: system prompt + requirement + few-shot examples."],
            ["OLLAMA_TIMEOUT",     "450",                   "HTTP timeout in seconds per LLM request."],
            ["OLLAMA_NUM_PREDICT", "3200",                  "Max tokens per response. Must fit a full 5-step TC JSON."],
        ],
        [2520, 1800, 5040],
    )

    add_heading(doc, "2.2  Embedding Settings", 2, BLUE)
    make_table(doc,
        ["Variable", "Default", "Purpose"],
        [
            ["EMBEDDING_BACKEND",      "ollama",                                     "Primary embedding backend. Options: ollama, sentence_transformer."],
            ["OLLAMA_EMBEDDING_MODEL", "qwen3-embedding:0.6b",                       "Ollama embedding model. Pulled automatically if missing."],
            ["OLLAMA_EMBEDDING_URL",   "http://localhost:11434/api/embeddings",       "Full URL for embedding API calls."],
            ["ST_EMBEDDING_MODEL",     "all-MiniLM-L6-v2",                           "SentenceTransformer fallback model if Ollama embedding unavailable."],
        ],
        [2520, 2520, 4320],
    )

    add_heading(doc, "2.3  Test Generation Settings", 2, BLUE)
    make_table(doc,
        ["Variable", "Default", "Purpose"],
        [
            ["MAX_TC_PER_REQ",   "6",        "Maximum test cases per requirement (positive + negative + edge case)."],
            ["INCLUDE_NEGATIVE", "true",      "Whether to generate negative test cases."],
            ["INCLUDE_EDGE",     "true",      "Whether to generate edge-case test cases."],
            ["TC_FORMAT",        "gherkin",   "Output format hint (gherkin = include Given/When/Then BDD fields)."],
            ["OUTPUT_DIR",       "./output",  "Directory where output files are written."],
            ["CHROMA_PERSIST_DIR", "./chroma_db", "Local filesystem path for persistent ChromaDB storage."],
        ],
        [2520, 1800, 5040],
    )

    doc.add_page_break()

    # ── Section 3: Data Models ────────────────────────────────────────────────
    add_heading(doc, "3. Data Models", 1, NAVY, bottom_border=HEX_TEAL)

    add_heading(doc, "3.1  Requirement", 2, BLUE)
    add_body(doc, "Produced by RequirementsReader. Passed to ChromaDB and TestCaseGenerator.")
    add_gap(doc, 60)
    make_table(doc,
        ["Field", "Type / Description"],
        [
            ["req_id",              "str — Unique identifier (e.g. REQ-001). Extracted from header or auto-assigned."],
            ["title",               "str — Short description of the requirement."],
            ["description",         "str — Full body text of the requirement."],
            ["priority",            "str — High / Medium / Low. Defaults to Medium if not found."],
            ["category",            "str — Domain category (e.g. authentication, onboarding)."],
            ["acceptance_criteria", "List[str] — Ordered list of AC statements parsed from the requirement text."],
            ["tags",                "List[str] — Metadata tags."],
            ["raw_text",            "str — Unprocessed original text for debug/audit purposes."],
        ],
        [2160, 7200],
    )

    add_heading(doc, "3.2  TestCase", 2, BLUE)
    add_body(doc, "The output unit of TestCaseGenerator. Exported by all three exporters.")
    add_gap(doc, 60)
    make_table(doc,
        ["Field", "Type / Description"],
        [
            ["tc_id",               "str — Auto-assigned sequential ID (TC-0001, TC-0002, …)."],
            ["req_id",              "str — Foreign key back to the originating Requirement."],
            ["title",               "str — One-line test case title."],
            ["description",         "str — One-sentence summary of what is being verified."],
            ["preconditions",       "str — Required system state and test data before execution."],
            ["test_type",           "str — positive | negative | edge_case."],
            ["priority",            "str — High | Medium | Low."],
            ["steps",               "List[TestStep] — Ordered list of {action, expected_result} pairs."],
            ["given / when / then", "str — BDD Gherkin scenario clauses."],
            ["component",           "str — UI screen name, dynamically resolved from domain vocab."],
            ["expected_outcome",    "str — Observable system state that proves the AC is satisfied."],
            ["tags",                "List[str] — Inherited + type-specific tags."],
            ["estimated_duration",  "str — Estimated execution time in minutes."],
        ],
        [2160, 7200],
    )

    doc.add_page_break()

    # ── Section 4: Component Reference ───────────────────────────────────────
    add_heading(doc, "4. Component Reference", 1, NAVY, bottom_border=HEX_TEAL)

    add_heading(doc, "4.1  LLM Client (llm_client.py)", 2, BLUE)
    add_body(doc, "Wraps the Ollama HTTP API. Manages structured JSON output, retry logic, and quality guards.")
    add_gap(doc, 60)

    add_heading(doc, "JSON Schema Enforcement", 3, ACCENT)
    add_body(doc, "Every generate_test_case() call passes a TESTCASE_JSON_SCHEMA to Ollama's format parameter, enforcing:")
    add_bullet(doc, "Required field list — all 13 TC fields must be present")
    add_bullet(doc, "minLength constraints on critical fields (title: 10, description: 20, expected_outcome: 25)")
    add_bullet(doc, "Enum constraints on test_type (positive | negative | edge_case) and priority")
    add_bullet(doc, "Stop sequences to prevent hallucinated text after the closing JSON brace")
    add_gap(doc, 80)

    add_heading(doc, "Retry Logic", 3, ACCENT)
    make_table(doc,
        ["Failure Type", "Action Taken"],
        [
            ["JSON parse error",                       "Retry with temperature +0.05 to increase response entropy."],
            ["Hollow response (empty required fields)", "Retry at same temperature with explicit field-fill instruction."],
            ["Both attempts fail",                     "Call synthesise_tc() — builds a deterministic fallback TC from the requirement text."],
        ],
        [2880, 6480],
    )

    add_heading(doc, "Sampling Parameters", 3, ACCENT)
    make_table(doc,
        ["Parameter", "Value / Reason"],
        [
            ["repeat_penalty", "1.15 — Strongly discourages repeated phrases; primary defence against hallucinated loops."],
            ["num_predict",    "3200 — Large enough to fit a 5-step TC JSON without mid-stream truncation."],
            ["top_k",          "40 — Explicit to prevent drift if Ollama defaults change between versions."],
            ["top_p",          "0.9 — Nucleus sampling cap. Balances diversity and coherence."],
        ],
        [2160, 7200],
    )

    add_heading(doc, "4.2  Requirements Reader (requirements_reader.py)", 2, BLUE)
    add_body(doc, "Detects file type by extension and routes to a format-specific parser.")
    make_table(doc,
        ["Format", "Parser Behaviour"],
        [
            [".csv / .xlsx", "Flexible column matching (id/req_id, title/name, description/details, acceptance_criteria/ac). Each row is one Requirement."],
            [".txt / .md",   "Splits on REQ-NNN headers, markdown H1/H2 headings, or numbered list items. Extracts AC bullet points."],
            [".pdf",         "Extracts text via PyPDF2 then applies same header/AC detection as .txt."],
            [".docx",        "Reads via python-docx paragraph by paragraph; same header/AC detection."],
            [".json",        "Expects array of objects with id/title/description/acceptance_criteria keys."],
        ],
        [1440, 7920],
    )

    add_heading(doc, "4.3  ChromaDB Vector Store (chroma_store.py)", 2, BLUE)
    make_table(doc,
        ["Collection", "Contents & Purpose"],
        [
            ["requirements", "Embedding of all requirement documents. Enables semantic similarity search for cross-requirement context during generation."],
            ["tc_examples",  "Approved test cases stored as few-shot examples. Filtered by ac_type and test_type. Persists across runs — the continuous learning store."],
            ["domain_vocab", "Extracted screen names and UI element names from requirement text. Enables dynamic component resolution without hardcoded maps."],
        ],
        [2160, 7200],
    )
    add_gap(doc, 60)

    add_heading(doc, "ChromaDB Initialisation Workaround (Python 3.14)", 3, ACCENT)
    add_body(doc,
        "ChromaDB's Settings class reads both os.environ AND the .env file via pydantic-settings "
        "(extra='forbid'). The _make_client() function works around this by:")
    add_bullet(doc, "Stripping all .env keys from os.environ temporarily")
    add_bullet(doc, "Changing CWD to a temp folder so pydantic-settings cannot find the .env file")
    add_bullet(doc, "Importing chromadb and pre-building a chromadb.Settings object while env is clean")
    add_bullet(doc, "Restoring CWD and env vars in the finally block")
    add_bullet(doc, "Passing the pre-built Settings to chromadb.Client() to skip the internal Settings() call")
    add_gap(doc, 80)

    add_heading(doc, "4.4  Test Generator (test_generator.py)", 2, BLUE)

    add_heading(doc, "AC Type Classifier", 3, ACCENT)
    make_table(doc,
        ["AC Type", "Detection Keywords"],
        [
            ["eligibility",  "can / cannot / eligible / ineligible / allowed / denied / age / days old"],
            ["ui_behaviour", "display / show / visible / screen / button / form / page / modal / prompt"],
            ["timing",       "seconds / minutes / within N / deadline / timeout / before / expires"],
            ["data_valid",   "validate / invalid / format / boundary / max / min / length / required field"],
            ["security",     "auth / password / OTP / biometric / token / encrypt / session / role"],
            ["general",      "Fallback for ACs not matching the above patterns."],
        ],
        [2160, 7200],
    )

    add_heading(doc, "Sanitiser Checks", 3, ACCENT)
    make_table(doc,
        ["Check", "Fix Applied"],
        [
            ["Repeated phrases (loop detection)",        "Scans for sentences appearing 3+ times; replaces entire field with synthesised value."],
            ["Snake_case leakage (variable names)",      "Strips tokens matching _SNAKE_RE (e.g. expected_result, test_type leaked from the schema)."],
            ["JSON / dict literals in text fields",      "Removes any {{ }} blocks that indicate the model echoed its own schema."],
            ["Hollow fields (empty or below min length)","Deterministically fills from requirement text rather than leaving a blank."],
        ],
        [2880, 6480],
    )

    doc.add_page_break()

    # ── Section 5: MCP Server ─────────────────────────────────────────────────
    add_heading(doc, "5. MCP Server & CLI Tools (mcp_server.py)", 1, NAVY, bottom_border=HEX_TEAL)

    add_heading(doc, "5.1  MCP Tools", 2, BLUE)
    make_table(doc,
        ["Tool", "Description"],
        [
            ["chromadb_stats()",                                     "Returns counts for all three ChromaDB collections."],
            ["list_examples(ac_type, test_type)",                    "Shows approved examples stored, with optional AC type / test type filter."],
            ["validate_test_case(tc_json)",                          "Quality check returning { passed, score (0-100), issues[] }."],
            ["repair_test_case(tc_json, field, issue_detail, ac)",   "Re-prompts the LLM to fix a single failing field. Surgical — does not regenerate the whole TC."],
            ["store_approved_tc(tc_json, ac_type, test_type, domain)","Validates then stores a TC as a few-shot example. Rejects TCs with validation errors."],
            ["retrieve_examples(ac_text, ac_type, test_type, n)",    "Returns the n closest approved TCs — shows what few-shot context the model will see."],
            ["domain_vocab_summary(category)",                       "Lists extracted screen names and UI element names from the domain_vocab collection."],
        ],
        [3240, 6120],
    )

    add_heading(doc, "5.2  Validation Scoring", 2, BLUE)
    make_table(doc,
        ["Penalty", "Trigger Condition"],
        [
            ["-15 points (error)",   "Title < 10 chars, description < 20 chars, fewer than 3 steps, generic step phrases ('perform the action'), trivial expected_outcome ('pass'), repeated sentences."],
            ["-5 points (warning)",  "Preconditions < 20 chars, any step missing expected_result, component is generic ('the application'), short Given/When/Then clauses."],
        ],
        [2160, 7200],
    )
    add_body(doc, "A TC passes if: score >= 75 AND no error-severity issues.")
    add_gap(doc, 80)

    add_heading(doc, "5.3  CLI Commands", 2, BLUE)
    make_table(doc,
        ["Command", "Usage"],
        [
            ["stats",         "python mcp_server.py stats"],
            ["validate-tc",   "python mcp_server.py validate-tc -i output/test_cases.csv --tc-id TC-0016"],
            ["repair-tc",     "python mcp_server.py repair-tc -i output/test_cases.csv --tc-id TC-0016 --field expected_outcome --ac \"<AC text>\""],
            ["store-example", "python mcp_server.py store-example -i output/test_cases.csv --tc-id TC-0002 --ac-type eligibility --test-type negative"],
            ["list-examples", "python mcp_server.py list-examples [--ac-type TYPE] [--test-type TYPE]"],
            ["vocab",         "python mcp_server.py vocab [--category CATEGORY]"],
        ],
        [2160, 7200],
    )

    doc.add_page_break()

    # ── Section 6: Orchestrator CLI ───────────────────────────────────────────
    add_heading(doc, "6. Orchestrator CLI (orchestrator.py)", 1, NAVY, bottom_border=HEX_TEAL)

    add_heading(doc, "6.1  CLI Flags", 2, BLUE)
    make_table(doc,
        ["Flag", "Default / Description"],
        [
            ["--requirements / -r", "REQUIRED. Path to requirements file."],
            ["--format / -f",       "both — Output: csv | zephyr | both."],
            ["--model / -m",        ".env value — Override OLLAMA_MODEL for this run only."],
            ["--max-tests",         ".env MAX_TC_PER_REQ — Override max TCs per requirement."],
            ["--output-dir / -o",   ".env OUTPUT_DIR — Override output directory."],
            ["--clear-db",          "Clear requirements + domain_vocab. Preserves tc_examples."],
            ["--clear-all",         "Wipe all three ChromaDB collections including approved examples."],
            ["--no-chroma",         "Skip ChromaDB entirely. Faster but no semantic context or few-shot examples."],
            ["--skip-llm-check",    "Skip the Ollama connection check at startup."],
        ],
        [2520, 6840],
    )

    add_heading(doc, "6.2  Example Commands", 2, BLUE)
    add_code(doc, "# Standard run")
    add_code(doc, "python orchestrator.py -r test_data/banking_rcd_requirements.csv")
    add_gap(doc, 40)
    add_code(doc, "# Clear index, re-run with different model")
    add_code(doc, "python orchestrator.py -r requirements.csv --clear-db -m qwen3:8b")
    add_gap(doc, 40)
    add_code(doc, "# Fast run without ChromaDB")
    add_code(doc, "python orchestrator.py -r requirements.csv --no-chroma --max-tests 3")
    add_gap(doc, 120)

    doc.add_page_break()

    # ── Section 7: Exporters ──────────────────────────────────────────────────
    add_heading(doc, "7. Exporters (exporters/exporters.py)", 1, NAVY, bottom_border=HEX_TEAL)

    add_heading(doc, "7.1  CSVExporter — output/test_cases.csv", 2, BLUE)
    add_body(doc, "Columns: TC ID, Requirement ID, Title, Description, Preconditions, Test Type, "
             "Priority, Category, Component, Steps, Expected Outcome, Given, When, Then, Tags, "
             "Estimated Duration, Generated At")
    add_body(doc, "Steps format:")
    add_code(doc, "1. Navigate to the Deposit section | Expected: RCD enrollment prompt appears")
    add_code(doc, "2. Tap 'Enroll Now' button          | Expected: Terms and Conditions screen loads")
    add_gap(doc, 80)

    add_heading(doc, "7.2  ZephyrScaleExporter — output/zephyr_test_cases.csv", 2, BLUE)
    make_table(doc,
        ["Zephyr Field", "Source"],
        [
            ["Name",                       "TestCase.title"],
            ["Status",                     "Always 'Draft' on initial import"],
            ["Priority",                   "Mapped from High/Medium/Low"],
            ["Component",                  "TestCase.component (dynamically resolved screen name)"],
            ["Test Script (Step-by-Step)", "Tab-separated: action \\t\\t expected_result per row"],
            ["Folder",                     "config.zephyr.folder_prefix / req_id"],
            ["Requirement",                "TestCase.req_id — links back to your Jira story"],
        ],
        [2880, 6480],
    )

    add_heading(doc, "7.3  JSONReportExporter — output/generation_report.json", 2, BLUE)
    add_bullet(doc, "generated_at (ISO timestamp), model_used")
    add_bullet(doc, "Summary: total_requirements, total_test_cases, avg_tests_per_requirement")
    add_bullet(doc, "breakdown_by_type: { positive: N, negative: N, edge_case: N }")
    add_bullet(doc, "breakdown_by_priority: { High: N, Medium: N, Low: N }")
    add_bullet(doc, "Full test_cases array — all fields for every generated TC")
    add_gap(doc, 120)

    doc.add_page_break()

    # ── Section 8: Quality Mechanisms ────────────────────────────────────────
    add_heading(doc, "8. Quality Mechanisms", 1, NAVY, bottom_border=HEX_TEAL)
    make_table(doc,
        ["Mechanism", "Where Applied"],
        [
            ["JSON Schema Enforcement",         "LLM: Ollama enforces field types, enum values, and minLength on every response."],
            ["2-Attempt Retry with Fallback",   "LLM: Parse failures and hollow responses are retried; hard failures use synthesise_tc()."],
            ["Repeat Penalty (1.15)",           "Sampling: Actively suppresses loop generation in every Ollama call."],
            ["Sanitiser (4 checks)",            "Post-LLM: Loop detection, snake_case removal, JSON leakage removal, hollow field fill."],
            ["AC Type-Specific Prompts",        "Prompt: Each of 6 AC types has tailored instructions for preconditions, steps, and outcomes."],
            ["Type-Specific Negative Framing",  "Prompt: Negative TCs use custom failure scenarios per AC type."],
            ["Few-Shot Example Injection",      "Prompt: Up to 2 approved TCs from ChromaDB are included in each generation prompt."],
            ["Validation Scoring (0-100)",      "Post-generation / CLI: Explicit issue enumeration with field-level granularity."],
            ["Surgical Field Repair",           "MCP tool: Failing fields are fixed with a targeted re-prompt, not a blind full-TC retry."],
            ["Human Feedback Loop",             "store-example CLI: Approved TCs become few-shot examples for future runs."],
        ],
        [3240, 6120],
    )

    doc.add_page_break()

    # ── Section 9: Dependencies ───────────────────────────────────────────────
    add_heading(doc, "9. Dependencies", 1, NAVY, bottom_border=HEX_TEAL)
    make_table(doc,
        ["Package", "Purpose"],
        [
            ["chromadb>=0.5.3",                    "Local vector store for requirements, examples, and domain vocab."],
            ["ollama>=0.1.8",                      "Python client for the local Ollama LLM server."],
            ["mcp>=1.0.0",                         "Model Context Protocol server (FastMCP) for Claude Desktop integration."],
            ["pydantic>=2.0.0 + pydantic-settings","Typed configuration and data validation."],
            ["python-dotenv>=1.0.0",               "Reads .env file into os.environ at startup."],
            ["rich>=13.0.0",                       "Coloured terminal output and progress display."],
            ["click>=8.1.0",                       "CLI argument parsing for orchestrator.py and mcp_server.py."],
            ["sentence-transformers>=2.2.2",       "Fallback embedding when Ollama embedding is unavailable."],
            ["python-docx>=1.1.0",                 "Reading .docx requirement files."],
            ["PyPDF2>=3.0.1",                      "Text extraction from .pdf requirement files."],
            ["pandas>=2.0.0 + openpyxl>=3.1.2",   "Reading .xlsx requirement files."],
        ],
        [3240, 6120],
    )

    doc.add_page_break()

    # ── Section 10: Setup ─────────────────────────────────────────────────────
    add_heading(doc, "10. Setup & Troubleshooting", 1, NAVY, bottom_border=HEX_TEAL)

    add_heading(doc, "10.1  Installation", 2, BLUE)
    add_code(doc, "pip install -r requirements.txt")
    add_code(doc, "ollama pull phi4-mini:3.8b")
    add_code(doc, "ollama pull qwen3-embedding:0.6b")
    add_body(doc, "Copy .env.example to .env and configure OLLAMA_BASE_URL and OLLAMA_MODEL.")
    add_gap(doc, 80)

    add_heading(doc, "10.2  Claude Desktop Integration", 2, BLUE)
    add_body(doc, "Add to claude_desktop_config.json:")
    add_code(doc, '{ "mcpServers": { "test-orchestrator": { "command": "python",')
    add_code(doc, '  "args": ["/path/to/mcp_server.py"], "env": { "OLLAMA_MODEL": "phi4-mini:3.8b" } } } }')
    add_gap(doc, 80)

    add_heading(doc, "10.3  Troubleshooting", 2, BLUE)
    make_table(doc,
        ["Error", "Fix"],
        [
            ["ChromaDB ValidationError: 17 extra inputs", "Already fixed in chroma_store.py via env isolation. Ensure latest file version."],
            ["Ollama connection refused",                  "Run: ollama serve. Ensure OLLAMA_BASE_URL matches the server address."],
            ["Model not found",                           "Run: ollama pull <model>. Orchestrator also attempts automatic pull."],
            ["TC validation score < 75",                  "Run validate-tc to see exact issues, then repair-tc for targeted field fixes."],
        ],
        [3240, 6120],
    )

    out_path = r"C:\Users\mohan\Downloads\files\Technical_Documentation.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# BUSINESS DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────
def build_business_doc():
    doc = Document()
    section = doc.sections[0]
    set_page_size(section)
    set_page_margins(section)

    style = doc.styles['Normal']
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ── Cover ─────────────────────────────────────────────────────────────────
    # Navy cover band
    cov_tbl = doc.add_table(rows=1, cols=1)
    cov_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(cov_tbl, 9360)
    cov_cell = cov_tbl.rows[0].cells[0]
    set_col_width(cov_cell, 9360)
    set_cell_bg(cov_cell, HEX_NAVY)
    set_cell_borders(cov_cell, HEX_NAVY, 1)
    set_cell_margins(cov_cell, 400, 400, 400, 400)

    tp = cov_cell.paragraphs[0]; tp.clear()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Twips(100)
    r = tp.add_run("AI TEST CASE GENERATOR")
    r.font.name = "Arial"; r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = WHITE

    sp2 = cov_cell.add_paragraph()
    sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp2.paragraph_format.space_after = Twips(80)
    r2 = sp2.add_run("Business Guide")
    r2.font.name = "Arial"; r2.font.size = Pt(17); r2.font.color.rgb = TEAL

    sp3 = cov_cell.add_paragraph()
    sp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = sp3.add_run("From Requirements to Executable Test Cases — Automatically")
    r3.font.name = "Arial"; r3.font.size = Pt(11); r3.font.italic = True
    r3.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)

    add_gap(doc, 120)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p3.add_run("March 2026  |  For: Product Owners, Business Analysts, QA Managers")
    r4.font.name = "Arial"; r4.font.size = Pt(10); r4.font.color.rgb = MID

    doc.add_page_break()

    # ── Executive Summary ─────────────────────────────────────────────────────
    add_heading(doc, "Executive Summary", 1, NAVY, bottom_border=HEX_TEAL)
    add_body(doc,
        "The AI Test Case Generator is an intelligent tool that reads your software requirement "
        "documents and automatically produces a complete, structured set of test cases — ready "
        "to execute or import into Zephyr Scale.")
    add_gap(doc, 60)
    add_body(doc, "It eliminates hours of manual test case writing by:")
    add_bullet(doc, "Reading requirements files in any common format (CSV, Excel, Word, PDF, plain text)")
    add_bullet(doc, "Understanding what each requirement needs to prove — positive scenarios, failure handling, and edge cases")
    add_bullet(doc, "Writing specific, step-by-step test cases with exact button names, input values, and expected outcomes")
    add_bullet(doc, "Exporting directly to Zephyr Scale or a standard spreadsheet format")
    add_bullet(doc, "Getting smarter over time as your team approves good examples and stores them as templates")
    add_gap(doc, 100)
    add_callout(doc,
        "Key Numbers from the Banking Sample Run",
        [
            "  10 requirements processed in a single run",
            "  60 test cases generated (6 per requirement: 2 positive, 2 negative, 2 edge case)",
            "  Every test case includes: title, description, preconditions, 5 executable steps, BDD scenario, expected outcome",
            "  Complete output ready for import into Zephyr Scale in under 10 minutes",
        ],
        fill=HEX_GREEN_BG,
        title_color=GREEN,
        border_color=HEX_GREEN,
    )

    doc.add_page_break()

    # ── The Problem ───────────────────────────────────────────────────────────
    add_heading(doc, "The Problem We Solve", 1, NAVY, bottom_border=HEX_TEAL)
    add_body(doc,
        "Writing test cases by hand is one of the most time-consuming activities in QA. "
        "A single requirement with 4-6 acceptance criteria can take 1-2 hours to convert into "
        "a thorough test case suite. Common pain points include:")
    add_gap(doc, 60)
    make_table(doc,
        ["Pain Point", "Business Impact"],
        [
            ["Writing the same boilerplate for every TC",  "QA engineers spend 40-60% of time on structure rather than thinking about test logic."],
            ["Missing negative and edge-case scenarios",   "Teams focus on the happy path and miss failure modes that appear in production."],
            ["Test cases too vague to execute",            "Steps like 'verify the login works' cannot be handed to a tester without interpretation."],
            ["Requirements change, tests are not updated", "Outdated test suites reduce confidence in release decisions."],
            ["Knowledge locked in individual engineers",   "New team members cannot quickly produce test cases at the expected quality level."],
        ],
        [3240, 6120],
    )

    doc.add_page_break()

    # ── How It Works ──────────────────────────────────────────────────────────
    add_heading(doc, "How It Works", 1, NAVY, bottom_border=HEX_TEAL)
    add_body(doc,
        "The tool runs a five-step pipeline that takes a requirements file as input and "
        "produces a ready-to-use test suite as output. The entire process runs locally — "
        "no data is sent to any external server.")
    add_gap(doc, 80)

    # Step flow as a table
    steps_tbl = doc.add_table(rows=2, cols=5)
    steps_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(steps_tbl, 9360)
    step_w = 1872
    step_data = [
        ("STEP 1", "Read Requirements", "Parses your requirements file and extracts each acceptance criterion"),
        ("STEP 2", "Build Context",     "Stores requirements in a local database for cross-reference"),
        ("STEP 3", "Generate",         "AI writes specific, executable test cases for each AC"),
        ("STEP 4", "Validate",         "Quality checks score every TC and flag weak content"),
        ("STEP 5", "Export",           "Outputs CSV and Zephyr Scale import files"),
    ]
    fills = [HEX_NAVY, HEX_BLUE, HEX_NAVY, HEX_BLUE, HEX_NAVY]
    for ci, (step_num, title, desc) in enumerate(step_data):
        label_cell = steps_tbl.rows[0].cells[ci]
        set_col_width(label_cell, step_w)
        set_cell_bg(label_cell, fills[ci])
        set_cell_borders(label_cell, HEX_WHITE, 4)
        set_cell_margins(label_cell, 80, 80, 100, 100)
        lp = label_cell.paragraphs[0]; lp.clear()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(step_num)
        lr.font.name = "Arial"; lr.font.size = Pt(9); lr.font.bold = True; lr.font.color.rgb = TEAL

        tp = label_cell.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(title)
        tr.font.name = "Arial"; tr.font.size = Pt(10); tr.font.bold = True; tr.font.color.rgb = WHITE

        desc_cell = steps_tbl.rows[1].cells[ci]
        set_col_width(desc_cell, step_w)
        set_cell_bg(desc_cell, HEX_LIGHT_BG)
        set_cell_borders(desc_cell, HEX_BLUE, 4)
        set_cell_margins(desc_cell, 80, 80, 100, 100)
        dp = desc_cell.paragraphs[0]; dp.clear()
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dr = dp.add_run(desc)
        dr.font.name = "Arial"; dr.font.size = Pt(9); dr.font.color.rgb = DARK

    doc.add_paragraph()
    add_gap(doc, 100)

    add_heading(doc, "Step 1: Read Requirements", 2, BLUE)
    add_body(doc, "The tool reads your requirements file and extracts each requirement's title, "
             "description, priority, and every acceptance criterion listed.")
    add_callout(doc,
        "Supported File Formats",
        ["  CSV / Excel (.csv, .xlsx) — structured requirement tables",
         "  Word documents (.docx)",
         "  PDF documents (.pdf)",
         "  Plain text or Markdown (.txt, .md)",
         "  JSON requirement exports"],
        fill=HEX_LIGHT_BG, title_color=BLUE, border_color=HEX_BLUE,
    )

    add_heading(doc, "Step 2: Build Context", 2, BLUE)
    add_body(doc,
        "Before generating test cases, the tool builds a knowledge base from your requirements. "
        "This allows it to find related requirements, learn screen names and UI element names, "
        "and retrieve previously approved test cases as quality templates. This knowledge base "
        "persists between runs.")
    add_gap(doc, 80)

    add_heading(doc, "Step 3: Generate Test Cases", 2, BLUE)
    add_body(doc, "For each acceptance criterion, the AI:")
    add_numbered(doc, "Classifies the criterion type (eligibility rule, UI behaviour, timing constraint, data validation, security control, or general)")
    add_numbered(doc, "Selects the right test approach for that type")
    add_numbered(doc, "Writes a positive test case (proving the happy path works)")
    add_numbered(doc, "Writes a negative test case (proving failure handling is correct)")
    add_numbered(doc, "Writes an edge-case test case (proving boundary conditions are handled)")
    add_gap(doc, 80)

    add_body(doc, "Every generated test case contains all of the following fields:")
    add_gap(doc, 40)
    make_table(doc,
        ["Field", "Example (Banking — RCD Eligibility)"],
        [
            ["Title",            "Verify ineligible user (account < 90 days) is shown specific rejection reason"],
            ["Description",      "Verifies that a customer with a new account (less than 90 days old) cannot enrol in RCD and sees the reason."],
            ["Preconditions",    "Customer is logged in with an account opened 45 days ago. Deposit section is accessible."],
            ["Test Steps",       "5 numbered steps, each naming the exact button or field, with an expected result after each action."],
            ["Expected Outcome", "Modal displays: 'Remote Cheque Deposit unavailable. Account must be at least 90 days old.' Enrol button is disabled."],
            ["BDD Scenario",     "Given / When / Then format for teams using Gherkin or Cucumber."],
            ["Component",        "Remote Cheque Deposit Enrollment Screen"],
            ["Priority",         "High"],
            ["Tags",             "eligibility, negative, rcd-enrollment"],
        ],
        [2160, 7200],
    )

    doc.add_page_break()

    # ── AC Types ──────────────────────────────────────────────────────────────
    add_heading(doc, "Understanding Acceptance Criterion Types", 1, NAVY, bottom_border=HEX_TEAL)
    add_body(doc,
        "The tool automatically identifies the type of each acceptance criterion and applies "
        "the most appropriate test strategy. You do not need to classify them manually.")
    add_gap(doc, 80)
    make_table(doc,
        ["AC Type", "What It Tests", "Banking Example"],
        [
            ["Eligibility",    "Who can and cannot use a feature. Rules about account age, customer status, or access rights.",        "Accounts less than 90 days old are shown a specific ineligibility reason and cannot enrol."],
            ["UI Behaviour",   "What the screen shows, how elements appear, and how the interface responds to user actions.",           "Terms and Conditions are displayed in full with scroll-to-bottom required before the Accept button is enabled."],
            ["Timing",         "Time limits, deadlines, and performance thresholds.",                                                   "OTP is delivered to the registered mobile number within 60 seconds of the request."],
            ["Data Validation","Field input rules, boundary values, and format requirements.",                                          "A cheque image rejected due to size shows an error message with the exact size limit."],
            ["Security",       "Authentication, authorisation, session management, and access controls.",                               "Biometric re-verification is attempted first; falls back to OTP only if biometric is unavailable."],
            ["General",        "Any acceptance criterion that doesn't fit the above categories.",                                       "The enrollment prompt appears on first navigation to the Deposit section only."],
        ],
        [1560, 2880, 4920],
    )

    doc.add_page_break()

    # ── Output Files ──────────────────────────────────────────────────────────
    add_heading(doc, "Output Files", 1, NAVY, bottom_border=HEX_TEAL)
    add_body(doc, "Every run produces three output files in the output/ directory:")
    add_gap(doc, 80)

    add_heading(doc, "1. Standard Test Cases (test_cases.csv)", 2, BLUE)
    add_body(doc,
        "A spreadsheet with one row per test case. Contains all fields: title, description, "
        "preconditions, steps, expected outcome, BDD scenario, priority, category, and component. "
        "Open in Excel or import into any test management tool.")
    add_callout(doc,
        "Steps Column Format",
        ["  Steps are written as numbered lines in a single cell:",
         "  1. Navigate to the Deposit section in the app  | Expected: Home screen > Deposit tab is displayed",
         "  2. Tap 'Remote Cheque Deposit'                 | Expected: RCD enrollment prompt appears",
         "  3. Tap 'Enroll Now'                            | Expected: Terms and Conditions screen loads"],
        fill=HEX_LIGHT_BG, title_color=BLUE, border_color=HEX_BLUE,
    )

    add_heading(doc, "2. Zephyr Scale Import (zephyr_test_cases.csv)", 2, BLUE)
    add_body(doc,
        "Ready to import directly into Zephyr Scale (Jira). Test cases are organised by "
        "requirement ID in folders (e.g., Project / REQ-001). The file includes the tab-separated "
        "step format that Zephyr Scale expects.")
    add_gap(doc, 60)
    make_table(doc,
        ["Zephyr Field", "Value Source"],
        [
            ["Name",        "Test case title"],
            ["Status",      "Draft (awaiting review)"],
            ["Priority",    "High / Medium / Low from the requirement"],
            ["Component",   "Screen name extracted from requirements (e.g., OTP Verification Screen)"],
            ["Labels",      "Tags from the requirement"],
            ["Folder",      "Project folder / Requirement ID"],
            ["Requirement", "The originating requirement ID — links back to your Jira story"],
        ],
        [2160, 7200],
    )

    add_heading(doc, "3. Generation Report (generation_report.json)", 2, BLUE)
    add_body(doc, "A machine-readable summary useful for CI/CD integration:")
    add_bullet(doc, "Total requirements processed and test cases generated")
    add_bullet(doc, "Breakdown by test type (positive / negative / edge case) and priority")
    add_bullet(doc, "Per-requirement counts — identify requirements with low coverage")
    add_bullet(doc, "Full model metadata (which AI model was used, generation timestamp)")

    doc.add_page_break()

    # ── Quality & Feedback Loop ───────────────────────────────────────────────
    add_heading(doc, "Quality Assurance & The Feedback Loop", 1, NAVY, bottom_border=HEX_TEAL)

    add_heading(doc, "Built-In Quality Checks", 2, BLUE)
    add_body(doc, "Every generated test case is automatically scored out of 100. The system checks for:")
    add_gap(doc, 60)
    make_table(doc,
        ["Quality Check", "What It Detects"],
        [
            ["Generic steps",          "Steps that say 'perform the action' instead of naming the actual button or field."],
            ["Trivial expected outcomes","Outcomes that say 'test passes' instead of describing the observable result."],
            ["Missing preconditions",  "Test cases with no system state or test data setup defined."],
            ["Repeated sentences",     "A sign the AI got confused and looped the same text — flagged for regeneration."],
            ["Vague components",       "Tests assigned to 'the application' instead of a specific screen name."],
        ],
        [2520, 6840],
    )
    add_body(doc,
        "A test case passes quality review when it scores 75/100 or above and has no error-level issues. "
        "Failing test cases can be repaired individually using the CLI tool — fixing only the specific "
        "field that failed, without regenerating the whole test case.")
    add_gap(doc, 100)

    add_heading(doc, "The Continuous Improvement Feedback Loop", 2, BLUE)
    add_body(doc,
        "The tool gets better with every run through a human-in-the-loop feedback mechanism:")
    add_gap(doc, 60)
    make_table(doc,
        ["Stage", "Who Acts", "What Happens"],
        [
            ["1. Generate", "Tool",         "Produces test cases using current examples and requirements."],
            ["2. Review",   "QA Engineer",  "Opens output/test_cases.csv and identifies the best test cases."],
            ["3. Approve",  "QA Engineer",  "Runs: python mcp_server.py store-example --tc-id TC-0002 --ac-type eligibility"],
            ["4. Learn",    "Tool",         "Stores the approved TC as a template in its local database."],
            ["5. Next Run", "Tool",         "Retrieves the approved template and uses it as an example when generating similar test cases."],
        ],
        [1440, 2160, 5760],
    )
    add_callout(doc,
        "Why This Matters",
        ["After 3-5 runs with your team approving good examples, the tool learns your team's style:",
         "  - The level of detail expected in step descriptions",
         "  - How your product names its screens and UI elements",
         "  - What a 'good' expected outcome looks like for your domain",
         "Output quality improves continuously without any code changes or AI retraining."],
        fill=HEX_GREEN_BG, title_color=GREEN, border_color=HEX_GREEN,
    )

    doc.add_page_break()

    # ── Business Benefits ─────────────────────────────────────────────────────
    add_heading(doc, "Business Benefits", 1, NAVY, bottom_border=HEX_TEAL)
    make_table(doc,
        ["Benefit", "Detail"],
        [
            ["10x faster test case creation",       "What takes a QA engineer 2+ hours per requirement takes the tool under 2 minutes. A 10-requirement module is covered in under 10 minutes."],
            ["Consistent structure every time",     "Every test case follows the same format: preconditions, 5 steps, expected outcome, BDD scenario. No variability between engineers or sprints."],
            ["Broader test coverage by default",    "The tool always generates positive, negative, and edge-case scenarios. Manual processes often skip negative cases under time pressure."],
            ["Specific, executable steps",          "The AI names the exact button label, field name, or UI element — not generic phrases. Tests can be handed to any tester without clarification."],
            ["Zero data privacy risk",              "Everything runs locally. No requirement text, AC content, or generated test cases leave your machine or network."],
            ["Improves with use",                   "Each approved test case becomes a quality template for future runs, continuously raising the bar without manual intervention."],
            ["Zephyr Scale ready",                  "Output is formatted for direct import — no reformatting or copy-paste required."],
            ["Works with your existing process",    "Reads the same requirements documents you already produce. No new tooling for business analysts."],
        ],
        [2880, 6480],
    )

    doc.add_page_break()

    # ── Glossary ───────────────────────────────────────────────────────────────
    add_heading(doc, "Glossary", 1, NAVY, bottom_border=HEX_TEAL)
    make_table(doc,
        ["Term", "Plain English Explanation"],
        [
            ["Acceptance Criterion (AC)", "A specific condition that a requirement must satisfy. For example: 'The OTP must be delivered within 60 seconds.'"],
            ["Test Case",                 "A documented scenario that checks whether one acceptance criterion is met. Includes preconditions, steps, and an expected outcome."],
            ["Positive Test Case",        "Tests the happy path — verifies that the system works correctly when everything is as expected."],
            ["Negative Test Case",        "Tests failure handling — verifies that the system handles an invalid or unexpected situation correctly."],
            ["Edge Case",                 "Tests boundary conditions — for example, a value exactly at the allowed limit (90 days old, not 89 or 91)."],
            ["BDD / Gherkin",             "A plain-language test format using Given / When / Then. Describes starting conditions, action, and expected result."],
            ["Zephyr Scale",              "A test management plugin for Jira used to organise, run, and report on test cases."],
            ["ChromaDB",                  "The local database used to store requirement context, approved test case templates, and screen/UI names."],
            ["Few-Shot Learning",         "The technique of including approved example test cases in the AI prompt so the model learns your quality standard."],
            ["Ollama",                    "The local AI server that runs the language model. No internet connection is required after initial model download."],
            ["MCP",                       "Model Context Protocol — a standard that allows Claude Desktop to use the tool's functions from within a conversation."],
        ],
        [2520, 6840],
    )

    out_path = r"C:\Users\mohan\Downloads\files\Business_Documentation.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    HEX_GREEN = "375623"
    print("Generating Technical Documentation...")
    build_technical_doc()
    print("Generating Business Documentation...")
    build_business_doc()
    print("Done.")
